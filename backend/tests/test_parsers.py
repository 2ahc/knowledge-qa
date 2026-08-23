from pathlib import Path

import pytest

from app.services.parsers import parse_document
from app.services.parsers.base import ParseError


def test_parse_txt_utf8_and_gbk(tmp_path: Path):
    p = tmp_path / "a.txt"
    p.write_text("第一段内容。\n第二段内容。", encoding="utf-8")
    segs = parse_document("txt", p)
    assert len(segs) == 1 and "第一段" in segs[0].text

    p2 = tmp_path / "b.txt"
    # long enough sample for reliable statistical detection
    p2.write_bytes(("GBK 编码的中文内容。" * 20).encode("gbk"))
    segs = parse_document("txt", p2)
    assert "GBK 编码" in segs[0].text


def test_parse_markdown_sections(tmp_path: Path):
    p = tmp_path / "doc.md"
    p.write_text("# 标题一\n内容A\n\n## 标题二\n内容B", encoding="utf-8")
    segs = parse_document("md", p)
    assert len(segs) == 2
    assert segs[0].meta.get("heading") == "标题一"
    assert segs[1].meta.get("heading") == "标题二"


def test_parse_docx(tmp_path: Path):
    from docx import Document

    doc = Document()
    doc.add_heading("产品介绍", level=1)
    doc.add_paragraph("这是第一段正文。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "名称"
    table.cell(0, 1).text = "价格"
    table.cell(1, 0).text = "奶茶"
    table.cell(1, 1).text = "12元"
    path = tmp_path / "doc.docx"
    doc.save(str(path))

    segs = parse_document("docx", path)
    all_text = "\n".join(s.text for s in segs)
    assert "产品介绍" in all_text
    assert "这是第一段正文。" in all_text
    assert "奶茶 | 12元" in all_text


def test_parse_xlsx(tmp_path: Path):
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "员工表"
    ws.append(["姓名", "部门"])
    ws.append(["张三", "研发"])
    ws.append([None, None])  # empty row skipped
    ws.append(["李四", "市场"])
    path = tmp_path / "data.xlsx"
    wb.save(str(path))

    segs = parse_document("xlsx", path)
    assert len(segs) == 1
    assert segs[0].meta["sheet"] == "员工表"
    assert "张三 | 研发" in segs[0].text


def test_parse_empty_file_raises(tmp_path: Path):
    p = tmp_path / "empty.txt"
    p.write_text("", encoding="utf-8")
    with pytest.raises(ParseError):
        parse_document("txt", p)


def test_unsupported_type_raises(tmp_path: Path):
    with pytest.raises(ParseError):
        parse_document("exe", tmp_path / "x")
