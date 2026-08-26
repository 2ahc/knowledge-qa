# Word(.docx) 解析：按标题分节提取段落，表格转成竖线分隔的文本。
from pathlib import Path

from docx import Document as DocxDocument

from app.services.parsers.base import ParseError, Segment


def parse(path: str | Path) -> list[Segment]:
    try:
        doc = DocxDocument(str(path))
    except Exception as e:
        raise ParseError(f"Word 文件无法解析: {e}")

    segments: list[Segment] = []
    current_heading = ""  # 当前所在章节标题（作为出处元信息）
    buffer: list[str] = []  # 当前章节累积的段落

    def flush():
        """把缓冲区内容输出为一个文本段（带上章节标题元信息）。"""
        text = "\n".join(buffer).strip()
        if text:
            meta = {"heading": current_heading} if current_heading else {}
            segments.append(Segment(text=text, meta=meta))
        buffer.clear()

    for para in doc.paragraphs:
        style = (para.style.name or "").lower() if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if style.startswith("heading"):
            # 遇到标题：结束上一节，开始新的一节
            flush()
            current_heading = text
            buffer.append(text)
        else:
            buffer.append(text)
            # 单节超过 2000 字就先落一段，避免单段过大
            if sum(len(s) for s in buffer) > 2000:
                flush()
    flush()

    # 表格处理：每行单元格用 " | " 拼接，让切片后仍保留行列结构
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append(" | ".join(cells))
        text = "\n".join(r for r in rows if r.strip(" |"))
        if text.strip():
            segments.append(Segment(text=text, meta={"type": "table"}))

    if not segments:
        raise ParseError("Word 文档中没有可提取的文本")
    return segments
